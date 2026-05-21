from __future__ import annotations

import re
import html
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


@dataclass
class ProposalData:
    company_name: str
    prompt: str
    html_content: str


DEFAULT_SELLER = {
    "name": "BrandProposAI RU",
    "person": "Менеджер отдела продаж",
    "address": "г. Киров, ул. Примерная, 1",
    "site": "www.brandproposai.ru",
    "phone": "+7 (8332) 555-900",
}


def _clean_text(value: str | None, default: str = "") -> str:
    value = (value or "").strip()
    return value if value else default


def extract_services(prompt: str) -> list[str]:
    """Простой генератор пунктов КП без внешних LLM: выделяет строки из запроса."""
    prompt = _clean_text(prompt, "Комплекс услуг по подготовке коммерческого предложения")
    raw_items: list[str] = []
    for line in re.split(r"[\n;]+", prompt):
        item = re.sub(r"^[-•\d.)\s]+", "", line).strip()
        if len(item) >= 3:
            raw_items.append(item)
    if not raw_items:
        raw_items = [prompt]
    result = []
    for item in raw_items[:6]:
        if not re.search(r"\b(услуг|поставка|разработка|дизайн|бренд|сайт|стенд|реклама)\b", item, flags=re.I):
            item = f"{item}"
        result.append(item[0].upper() + item[1:])
    return result


def generate_proposal_html(company_name: str, prompt: str) -> str:
    company_name = _clean_text(company_name, "Клиент")
    services = extract_services(prompt)
    lis = "".join(f"<li>{html.escape(service)}</li>" for service in services)
    return f"""
        <p style="text-align:center;"><strong>{html.escape(company_name)}</strong></p>
        <h2>Коммерческое предложение</h2>
        <p><strong>Уважаемые партнёры!</strong></p>
        <p>Компания «BrandProposAI RU» предлагает рассмотреть индивидуальное коммерческое предложение, подготовленное с учётом указанного запроса и требуемого состава работ.</p>
        <p>Предлагаем выполнить следующие работы и услуги:</p>
        <ul>{lis}</ul>
        <p>Состав работ может быть уточнён после согласования технического задания, сроков выполнения и дополнительных требований заказчика.</p>
        <p><strong>Предварительная стоимость:</strong> рассчитывается индивидуально после подтверждения объёма работ.</p>
        <p><strong>Преимущества предложения:</strong></p>
        <ul>
            <li>персонализация текста под задачи клиента;</li>
            <li>единый стиль исходящих коммерческих предложений;</li>
            <li>возможность редактирования документа перед сохранением;</li>
            <li>экспорт результата в DOCX и PDF.</li>
        </ul>
        <p>На общую сумму: по согласованию сторон.</p>
        <p style="margin-top:2rem;">С уважением,</p>
        <p><strong>{DEFAULT_SELLER['person']}</strong><br>{DEFAULT_SELLER['name']}<br>{DEFAULT_SELLER['address']}<br>{DEFAULT_SELLER['site']}<br>тел. {DEFAULT_SELLER['phone']}</p>
    """.strip()


def html_to_plain_lines(html_content: str) -> list[str]:
    soup = BeautifulSoup(html_content or "", "html.parser")
    lines: list[str] = []
    for node in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = " ".join(node.get_text(" ", strip=True).split())
        if not text:
            continue
        if node.name == "li":
            text = f"- {text}"
        lines.append(text)
    if not lines:
        text = " ".join(BeautifulSoup(html_content or "", "html.parser").get_text(" ", strip=True).split())
        if text:
            lines.append(text)
    return lines


def create_docx(path: Path, data: ProposalData) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    number = doc.add_paragraph("№ _________")
    number.alignment = WD_ALIGN_PARAGRAPH.LEFT

    recipient = doc.add_paragraph(_clean_text(data.company_name, "Получатель"))
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recipient.runs[0].bold = True

    title = doc.add_paragraph("Коммерческое предложение")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    for line in html_to_plain_lines(data.html_content):
        if line.lower() == "коммерческое предложение":
            continue
        p = doc.add_paragraph()
        if line.startswith("- "):
            p.style = doc.styles["List Bullet"]
            p.add_run(line[2:])
        else:
            p.add_run(line)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("")
    footer = doc.add_paragraph(f"Документ сформирован информационной системой BrandProposAI RU: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(path)
    return path


def _register_pdf_font() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
    ]
    for font_path in candidates:
        if Path(font_path).exists():
            name = Path(font_path).stem
            pdfmetrics.registerFont(TTFont(name, font_path))
            return name
    return "Helvetica"


def create_pdf(path: Path, data: ProposalData) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    font = _register_pdf_font()
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=18*mm, bottomMargin=18*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="RuNormal", fontName=font, fontSize=11, leading=16, spaceAfter=8))
    styles.add(ParagraphStyle(name="RuTitle", fontName=font, fontSize=15, leading=20, alignment=1, spaceAfter=12))
    story = [Paragraph(html.escape(_clean_text(data.company_name, "Получатель")), styles["RuNormal"]), Paragraph("Коммерческое предложение", styles["RuTitle"])]
    for line in html_to_plain_lines(data.html_content):
        if line.lower() == "коммерческое предложение":
            continue
        story.append(Paragraph(html.escape(line), styles["RuNormal"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Документ сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles["RuNormal"]))
    doc.build(story)
    return path
